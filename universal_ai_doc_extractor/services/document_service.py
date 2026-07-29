"""Document service - orchestrates PDF, OCR, and text extraction."""

import logging
import time
from pathlib import Path
from typing import Optional

from config.constants import SUPPORTED_IMAGE_FORMATS, SUPPORTED_DOC_FORMATS
from services.pdf_service import PDFService
from services.ocr_service import OCRService
from database.repositories import DocumentRepository, ActivityLogRepository
from models.enums import LogCategory
from utils.helpers import is_image_file, is_pdf_file


logger = logging.getLogger(__name__)


class DocumentService:
    def __init__(self) -> None:
        self.pdf_service = PDFService()
        self.ocr_service = OCRService()
        self.doc_repo = DocumentRepository()
        self.log_repo = ActivityLogRepository()

    def extract_text(self, file_path: str, password: Optional[str] = None,
                     language: str = "eng") -> str:
        start = time.time()
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        text = ""
        if is_pdf_file(file_path):
            text = self._extract_pdf(file_path, password)
        elif is_image_file(file_path):
            text = self._extract_image(file_path, language)
        elif file_path.lower().endswith(".txt"):
            text = path.read_text(encoding="utf-8", errors="replace")
        elif file_path.lower().endswith(".csv"):
            text = path.read_text(encoding="utf-8", errors="replace")
        elif file_path.lower().endswith(".docx"):
            text = self._extract_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_path}")

        elapsed_ms = int((time.time() - start) * 1000)
        self.log_repo.log(
            "document_extraction", LogCategory.DOCUMENT.value,
            f"Extracted {len(text)} chars from {path.name} in {elapsed_ms}ms",
            elapsed_ms,
        )

        return text

    def _extract_pdf(self, file_path: str, password: Optional[str] = None) -> str:
        try:
            return self.pdf_service.extract_text(file_path, password)
        except ValueError as e:
            if "password" in str(e).lower():
                raise
            logger.info("PDF text extraction failed, trying OCR: %s", e)
            return self._ocr_pdf(file_path)
        except Exception as e:
            logger.info("PDF text extraction error, trying OCR: %s", e)
            return self._ocr_pdf(file_path)

    def _extract_image(self, file_path: str, language: str = "eng") -> str:
        processed = self.ocr_service.preprocess_image(file_path)
        return self.ocr_service.extract_text(processed, language)

    def _ocr_pdf(self, file_path: str, language: str = "eng") -> str:
        try:
            import pypdfium2 as pdfium
            pdf = pdfium.PdfDocument(file_path)
            texts = []
            for i in range(len(pdf)):
                page = pdf[i]
                bitmap = page.render(scale=2)
                img_bytes = bitmap.encode("png")

                import tempfile
                with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                    tmp.write(img_bytes)
                    tmp_path = tmp.name

                try:
                    page_text = self.ocr_service.extract_text(tmp_path, language)
                    texts.append(f"--- Page {i + 1} ---\n{page_text}")
                finally:
                    Path(tmp_path).unlink(missing_ok=True)

            pdf.close()
            return "\n\n".join(texts)
        except ImportError:
            raise
        except Exception as e:
            logger.error("PDF OCR failed: %s", e)
            raise

    def _extract_docx(self, file_path: str) -> str:
        try:
            from docx import Document
            doc = Document(file_path)
            return "\n".join(p.text for p in doc.paragraphs)
        except ImportError:
            logger.warning("python-docx not installed, trying basic extraction")
            raise
        except Exception as e:
            logger.error("DOCX extraction failed: %s", e)
            raise
